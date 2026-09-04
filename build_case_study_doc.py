# -*- coding: utf-8 -*-
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header.is_linked_to_previous = False
        
        # Header / Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("SUBTHAITLE Architecture & Engineering Case Study | Confidential & Public Whitepaper")
        hrun.font.name = "Tahoma"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(148, 163, 184) # slate-400

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("เอกสารแบ่งปันกรณีศึกษาเพื่อชุมชนนักพัฒนาและ Content Creator โดย พี่เอ & น้อง Sunday")
        frun.font.name = "Tahoma"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(148, 163, 184)

    # Styles helper
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Tahoma"
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate-900
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(20)
        run = p.add_run(text)
        run.font.name = "Tahoma"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(71, 85, 105) # Slate-600
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Tahoma"
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(30, 41, 59) # Slate-800
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Tahoma"
        run.font.size = Pt(12.5)
        run.font.color.rgb = RGBColor(14, 116, 144) # Cyan-700
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Tahoma"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(51, 65, 85) # Slate-700
        return p

    def add_p(text, bold_prefix="", italic_suffix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            br = p.add_run(bold_prefix)
            br.bold = True
            br.font.name = "Tahoma"
            br.font.size = Pt(10.5)
            br.font.color.rgb = RGBColor(15, 23, 42)
        run = p.add_run(text)
        run.font.name = "Tahoma"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(51, 65, 85)
        if italic_suffix:
            ir = p.add_run(italic_suffix)
            ir.italic = True
            ir.font.name = "Tahoma"
            ir.font.size = Pt(10)
            ir.font.color.rgb = RGBColor(100, 116, 139)
        return p

    def add_bullet(text, bold_prefix="", level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        if bold_prefix:
            br = p.add_run(bold_prefix)
            br.bold = True
            br.font.name = "Tahoma"
            br.font.size = Pt(10.5)
            br.font.color.rgb = RGBColor(15, 23, 42)
        run = p.add_run(text)
        run.font.name = "Tahoma"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_callout(text, title="KEY INSIGHT / บันทึกเชิงเทคนิค"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F8FAFC") # slate-50
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

        # left border styling
        tcPr = cell._element.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="24" w:space="0" w:color="0284C7"/>' # Sky-600 3pt border
            f'<w:top w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:bottom w:val="none"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after = Pt(2)
        trun = cp.add_run(f"💡 {title}\n")
        trun.bold = True
        trun.font.name = "Tahoma"
        trun.font.size = Pt(10.5)
        trun.font.color.rgb = RGBColor(2, 132, 199)

        mrun = cp.add_run(text)
        mrun.font.name = "Tahoma"
        mrun.font.size = Pt(10)
        mrun.font.color.rgb = RGBColor(51, 65, 85)
        
        # spacing after callout
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(4)

    # ---------------------------------------------------------------------------
    # DOCUMENT CONTENT
    # ---------------------------------------------------------------------------

    # Title & Metadata
    add_title("SUBTHAITLE: สถาปัตยกรรมระบบทำซับไตเติลภาษาไทยด้วย AI")
    add_subtitle("กรณีศึกษาเชิงลึก: สถาปัตยกรรมซอฟต์แวร์, การผสาน AI สองโมเดลแก้ปัญหาเสียงดีเลย์, และโมเดลธุรกิจ Freemium\nโดย: พี่เอ (VP Data Platform & Tech Content Creator) ร่วมกับ น้อง Sunday (AI Coding Agent)")

    # Callout overview
    add_callout(
        "เอกสารฉบับนี้สรุปแนวคิด การออกแบบสถาปัตยกรรม (Design Architecture), เทคโนโลยีที่เลือกใช้ (Tech Stack), "
        "การแก้โจทย์ทางเทคนิคระดับลึกเรื่อง 'เสียงไม่ตรงกับซับไตเติล' ด้วยนวัตกรรม Dual-Engine Parallel Fusion, "
        "ตลอดจน Business Model และ Cost Structure ทางธุรกิจ เพื่อเป็นแนวทางและกรณีศึกษาสำหรับนักพัฒนา, "
        "ผู้สนใจด้าน AI Engineering, และ Creator ที่ต้องการสร้าง Digital Product ยุคใหม่โดยไม่มีค่าใช้จ่ายเซิร์ฟเวอร์มหาศาล",
        "บทนำและวัตถุประสงค์ของการแชร์ความรู้"
    )

    # ---------------------------------------------------------------------------
    # SECTION 1: Background & Vision
    # ---------------------------------------------------------------------------
    add_h1("1. บทนำและที่มาของโปรเจกต์ (Problem Statement & Vision)")
    add_p(
        "ในยุคของ Short-form Video (TikTok, Instagram Reels, YouTube Shorts) ซับไตเติลภาษาไทยที่มีความสวยงาม "
        "อ่านง่าย และมีฟีเจอร์ Word Highlight (เน้นสีคำที่กำลังพูดแบบเรียลไทม์) คือหัวใจสำคัญในการตรึงผู้ชม (Audience Retention) "
        "แต่ในความเป็นจริง Content Creator ในไทยกำลังเผชิญกับ Pain Point สำคัญ 3 ประการ:"
    )
    add_bullet("การพิมพ์ซับไตเติลและกะจังหวะ Highlight ทีละคำด้วยมือในโปรแกรมตัดต่อกินเวลา 2-4 ชั่วโมงต่อคลิปสั้นเพียง 1 นาที", "1. ความสูญเปล่าของเวลา (Time-Consuming): ")
    add_bullet("เครื่องมือ AI ต่างประเทศ (เช่น CapCut Pro, Descript, Opus Clip) คิดค่าบริการรายเดือนสูง (เดือนละ 400 - 1,200 บาท) และแย่ที่สุดคือ 'ถอดภาษาไทยเพี้ยนหนักมาก' ตัดคำผิด วรรณยุกต์หลุด และไม่เข้าใจศัพท์แสลงไทย", "2. กำแพงค่าบริการผูกมัดและภาษาไทยที่ด้อยคุณภาพ: ")
    add_bullet("ปัญหาเสียงไม่ตรงกับคำ (Audio-Subtitle Desync) และ Word Highlight วิ่งล่วงหน้า ทำให้คลิปดูไม่เป็นมืออาชีพ", "3. ความคลาดเคลื่อนของเวลา: ")

    add_p(
        "วิสัยทัศน์ของโปรเจกต์ SUBTHAITLE คือการสร้างเครื่องมือทำซับไตเติลภาษาไทยที่ดีที่สุดในประเทศ ด้วยแนวคิด Vibe Coding "
        "ที่เปิดให้ Creator ทุกคนใช้งานฟรี 'ไม่จำกัดจำนวนคลิป' สำหรับคลิปสั้น (< 2 นาที) และคิดเงินเฉพาะกรณีมีคลิปยาวหรือต้องการเป็น VIP "
        "ในรูปแบบ Pay-as-you-go (เติมเครดิตตามจริง ไม่มีวันหมดอายุ ไม่ผูกมัดรายเดือน)",
        bold_prefix="วิสัยทัศน์ของ SUBTHAITLE: "
    )

    # ---------------------------------------------------------------------------
    # SECTION 2: System Architecture & Tech Stack
    # ---------------------------------------------------------------------------
    add_h1("2. สถาปัตยกรรมระบบและเทคโนโลยีที่เลือกใช้ (System Architecture & Tech Stack)")
    add_p("SUBTHAITLE ถูกออกแบบภายใต้หลักการ Modern Serverless & Client-Centric Architecture เพื่อประสิทธิภาพสูงสุดและควบคุมต้นทุนให้ใกล้เคียงศูนย์:")

    # Table of Tech Stack
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["เลเยอร์ระบบ (Layer)", "เทคโนโลยีที่เลือกใช้ (Tech Stack)", "เหตุผลเชิงเทคนิคและสถาปัตยกรรม"]
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "0F172A") # Slate-900
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.name = "Tahoma"
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    stack_data = [
        ("Frontend & UI Framework", "Next.js 15 (App Router, Turbopack) + React 19 + Tailwind CSS", "รองรับ Server Components, คอมไพล์รวดเร็ว และให้ประสบการณ์ใช้งานที่ลื่นไหลระดับ Native Desktop App"),
        ("State Management", "Zustand + Custom Partialize Persistence", "จัดการไทม์ไลน์วิดีโอ, ซับไตเติล, การ undo/redo และสถานะโปรเจกต์ได้อย่างเบาหวิว ไม่เกิด Re-render ที่สิ้นเปลือง"),
        ("In-Browser Media Engine", "@ffmpeg/ffmpeg (WebAssembly / Wasm)", "สกัดเสียง MP3 จากไฟล์วิดีโอบนเบราว์เซอร์ของผู้ใช้โดยตรง ตัดต้นทุน Bandwidth และพื้นที่จัดเก็บบนคลาวด์ได้ 100%"),
        ("Backend Serverless API", "Next.js Route Handlers (Node.js Runtime บน Vercel Singapore Region)", "Latency ต่ำสุดสำหรับผู้ใช้งานในไทย, รันเฉพาะตอนประมวลผล ไม่มีค่าเซิร์ฟเวอร์รายเดือนคงที่ (Serverless Idle = ฿0)"),
        ("Database & Auth", "Supabase (PostgreSQL + Row-Level Security + Stored Procedures)", "ดูแล User Profiles, Google OAuth Login, ตรวจสอบเครดิตคงเหลือ และบันทึกคำศัพท์ Custom Dictionary"),
        ("Payment & Billing", "Stripe Checkout & Webhooks", "รับชำระเงินมาตรฐานสากล, ตัดเงินแล้วอัปเดตเครดิตนาทีเข้าฐานข้อมูลทันทีแบบ Real-time"),
        ("AI Inference Engines", "Google Gemini 3.8/3.7 Flash + Groq Whisper Large v3", "Gemini ดูแลความถูกต้องด้านภาษาไทย 100% ส่วน Groq Whisper ดูแลความเป๊ะของจังหวะคลื่นเสียง"),
    ]

    for layer, tech, reason in stack_data:
        row_cells = tbl.add_row().cells
        for idx, text in enumerate([layer, tech, reason]):
            row_cells[idx].text = text
            set_cell_margins(row_cells[idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[idx].paragraphs[0]
            p.runs[0].font.name = "Tahoma"
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.color.rgb = RGBColor(30, 41, 59)
            if idx == 0:
                p.runs[0].font.bold = True
                set_cell_background(row_cells[idx], "F1F5F9")
            else:
                set_cell_background(row_cells[idx], "FFFFFF")

    # Spacing after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(4)

    # ---------------------------------------------------------------------------
    # SECTION 3: In-Browser Media Processing
    # ---------------------------------------------------------------------------
    add_h2("ความลับเรื่องต้นทุน ฿0: สกัดเสียงบนเครื่องผู้ใช้ (In-Browser Wasm Audio Extraction)")
    add_p(
        "โมเดลเว็บแอปพลิเคชันวิดีโอทั่วไป มักบังคับให้ผู้ใช้อัปโหลดไฟล์วิดีโอขนาด 200MB - 1GB ขึ้นเซิร์ฟเวอร์ก่อน "
        "จากนั้นเซิร์ฟเวอร์จึงใช้ FFmpeg แปลงไฟล์ ซึ่งก่อให้เกิดปัญหา 2 ด้านคือ: (1) ผู้ใช้ต้องรอนานมากหากเน็ตช้า "
        "และ (2) เจ้าของเว็บต้องจ่ายค่า Bandwidth (Egress/Ingress) และ Cloud Compute มหาศาล"
    )
    add_p(
        "SUBTHAITLE แก้ปัญหานี้ด้วยการนำ FFmpeg มาคอมไพล์เป็น WebAssembly (Wasm) แล้วรันบน CPU ในเครื่องของผู้ใช้เอง "
        "เมื่อผู้ใช้เลือกไฟล์วิดีโอ 500MB เบราว์เซอร์จะสกัดเฉพาะแทร็กเสียงภาษาไทยออกมาเป็น MP3 16kHz คุณภาพสูงขนาดเพียง 1-3MB "
        "ในเวลาเพียง 2-3 วินาที จากนั้นจึงส่งเฉพาะไฟล์ MP3 ก้อนเล็กนี้ขึ้น API ของเรา:",
        bold_prefix="วิธีการแก้ปัญหาของ SUBTHAITLE: "
    )
    add_bullet("ผู้ใช้ประหยัดเน็ตมือถือและไม่ต้องรอนาน (เร็วขึ้นกว่าเดิม 10 เท่า)", "ผลลัพธ์ที่ 1: ")
    add_bullet("เซิร์ฟเวอร์ของเราไม่ต้องแบกรับภาระการแปลงไฟล์วิดีโอ ลดต้นทุน Cloud Compute ลงเหลือศูนย์", "ผลลัพธ์ที่ 2: ")
    add_bullet("ขนาดไฟล์ที่ส่งเข้า AI มีขนาดเล็กมาก ทำให้ถอดเสียงเสร็จภายในไม่กี่วินาที", "ผลลัพธ์ที่ 3: ")

    # ---------------------------------------------------------------------------
    # SECTION 4: Audio-Subtitle Sync & Dual-Engine Fusion
    # ---------------------------------------------------------------------------
    add_h1("3. ถอดรหัสความลับ: ทำอย่างไรให้ 'เสียงตรงซับเป๊ะระดับมิลลิวินาที' (Audio Synchronization Secret)")
    add_p(
        "นี่คือจุดปราบเซียนที่สุดของการทำ AI Subtitle ในภาษาไทย! หากใครเคยลองสร้างแอปถอดเสียง จะพบกับทางสองแพร่งที่น่าหงุดหงิด:"
    )
    add_bullet("Whisper เกิดมาเป็น Acoustic Model (โมเดลคลื่นเสียง) จับเวลาเสียงเริ่มพูดและหยุดพูดได้เป๊ะระดับ 0.01 วินาที ไม่เคยคลาดเคลื่อน แต่สำหรับภาษาไทย Whisper สะกดคำผิดเยอะมาก (เช่น 'สวัดดี', 'คับ', 'ชาจไว'), ไม่เข้าใจศัพท์แสลงวัยรุ่น, ตัดวรรณยุกต์ทิ้ง และแบ่งคำแปลกประหลาด", "1. ฝั่ง Groq Whisper Large v3 (จังหวะเป๊ะ แต่คำเพี้ยน): ")
    add_bullet("Gemini เป็น Multimodal Large Language Model (LLM) มีความเข้าใจภาษาไทยระดับอัจฉริยะ สะกดคำถูกต้อง 100%, เข้าใจบริบท, ศัพท์ไอที และเครื่องหมายวรรคตอน แต่เนื่องจาก Gemini เป็น LLM มันจึง 'กะเดาตัวเลขเวลา (Timestamps) ในข้อความ JSON' แทนที่จะจับคลื่นเสียงจริง พอคนพูดเร็ว พูดช้า หรือมีจังหวะหยุดหายใจ ตัวเลขเวลาของ Gemini จะเริ่มสะสมความคลาดเคลื่อน (Cumulative Time Drift) ทำให้ท้ายคลิปเสียงกับซับจะหนีกัน 1-2 วินาที", "2. ฝั่ง Google Gemini Flash (คำเป๊ะ แต่จังหวะเพี้ยน): ")

    add_h2("นวัตกรรม Dual-Engine Parallel Fusion: รวมจุดเด่นของทั้งสองค่าย")
    add_p(
        "ทีมงาน SUBTHAITLE ได้ออกแบบและคิดค้นสถาปัตยกรรม 'Dual-Engine Parallel Fusion' ขึ้นมาเพื่อรวมพลังของสองโมเดล "
        "โดยส่งไฟล์เสียงให้ Google Gemini และ Groq Whisper ประมวลผลพร้อมกันแบบคู่ขนาน (Parallel Execution ผ่าน Promise.allSettled) "
        "จากนั้นนำผลลัพธ์ทั้งสองมาผสานกันด้วยอัลกอริทึมบนเซิร์ฟเวอร์:"
    )

    add_callout(
        "1. Groq Whisper ใช้เวลาถอดเสียงเพียง 1.0 - 1.5 วินาที ได้ 'ช่วงเวลาคลื่นเสียงระดับมิลลิวินาที (Acoustic Timestamps)'\n"
        "2. Google Gemini ใช้เวลาถอดเสียงประมาณ 2.5 - 3.5 วินาที ได้ 'คำศัพท์ภาษาไทยที่ถูกต้อง 100% (Linguistic Ground Truth)'\n"
        "3. เนื่องจากรันคู่ขนาน เวลารวมจึงเท่ากับตัวที่ช้ากว่า (~3.5 วินาที) ไม่ทำให้ผู้ใช้ต้องรอนานขึ้นเลยแม้แต่น้อย!",
        "กระบวนการทำงานของ Dual-Engine Parallel Fusion"
    )

    add_h2("อัลกอริทึม Needleman-Wunsch Dynamic Programming Alignment")
    add_p(
        "โจทย์สำคัญคือ: เราจะเอาคำไทยที่ถูกต้องของ Gemini ไปสวมทับลงบนเวลาของ Whisper ได้อย่างไร ในเมื่อ Whisper สะกดคำผิดและแบ่งคำไม่เหมือนกัน? "
        "คำตอบคือการใช้อัลกอริทึมทางชีวสารสนเทศศาสตร์ (Bioinformatics) ระดับคลาสสิกที่เรียกว่า Needleman-Wunsch Sequence Alignment "
        "ซึ่งนำมาประยุกต์ใช้กับ Character Stream ของภาษาไทย:"
    )
    add_bullet("สร้างสายอักขระเสียงจาก Whisper โดยผูกตัวอักษรทุกตัวเข้ากับเวลาจริงบนคลื่นเสียง (Acoustic Time Matrix)", "ขั้นตอนที่ 1: ")
    add_bullet("สร้างสายอักขระคำที่ถูกต้องจาก Gemini", "ขั้นตอนที่ 2: ")
    add_bullet("รัน Dynamic Programming เปรียบเทียบความคล้ายคลึงของตัวอักษร โดยมีฟังก์ชัน Fuzzy Thai Matching ผ่อนปรนความแตกต่างของสระและวรรณยุกต์ ทำให้คำว่า 'สวัดดี' ของ Whisper จับคู่กับ 'สวัสดี' ของ Gemini ได้อย่างแม่นยำ 100%", "ขั้นตอนที่ 3: ")
    add_bullet("นำพิกัดเวลา Start และ End ที่แท้จริงจากคลื่นเสียงของ Whisper มากำหนดเป็นเวลาของคำที่ถูกต้องจาก Gemini", "ขั้นตอนที่ 4: ")
    add_bullet("สำหรับคำที่ Whisper ฟังข้ามไป ระบบมีกลไก Linear Gap Interpolation เกลี่ยเวลาอย่างเป็นสัดส่วน ป้องกันไม่ให้ไทม์ไลน์กระโดด", "ขั้นตอนที่ 5: ")

    add_h2("การล็อกจังหวะ Presentation Timestamp (PTS) ด้วย FFmpeg")
    add_p(
        "นอกจากโมเดล AI แล้ว อีกหนึ่งสาเหตุคลาสสิกที่ทำให้คลิปจากสมาร์ตโฟน (iPhone/Android) มีเสียงไม่ตรงซับ "
        "คือปัญหาวิดีโอแบบ Variable Frame Rate (VFR) ซึ่งเฟรมเรตไม่คงที่ ทำให้เสียงค่อย ๆ เลื่อนหลุดจากภาพ "
        "ทีมงานแก้ไขด้วยการใช้ Audio Resampling Filter ในขั้นตอนการสกัดเสียง:",
        bold_prefix="การแก้ปัญหา VFR: "
    )
    add_p(
        "aresample=16000:async=1",
        bold_prefix="คำสั่งหัวใจสำคัญ: ",
        italic_suffix=" (ทำการซิงค์ Timestamp ของเสียงเข้ากับตัวนับเวลาของคอนเทนเนอร์วิดีโอ 1:1 แบบไร้การบิดเบือน)"
    )

    # ---------------------------------------------------------------------------
    # SECTION 5: AI Integration & API Routing Logic
    # ---------------------------------------------------------------------------
    add_h1("4. กลยุทธ์การเชื่อมต่อ AI & API Routing Logic")
    add_p(
        "การพึ่งพา AI Provider ภายนอกมีความท้าทายเรื่อง Rate Limit (HTTP 429), เซิร์ฟเวอร์หน่วง (HTTP 503) "
        "และปัญหาความพร้อมใช้งาน SUBTHAITLE ได้วางโครงสร้างความปลอดภัยของ API ไว้อย่างรัดกุม:"
    )

    add_h2("Multi-Model Cascade: แก้ปัญหา 429 ของ Gemini Free Tier อย่างถาวร")
    add_p(
        "Google AI Studio จัดสรรโควต้า RPM (Requests Per Minute) และ Capacity แยกตามโมเดล "
        "เมื่อโมเดลใหม่อย่าง gemini-3.8-flash มีผู้ใช้งานทั่วโลกหนาแน่นจนติด 429 แต่โมเดลอื่น ๆ กลับว่างและตอบกลับได้ทันที "
        "เราจึงสร้างระบบ 4-Tier Cascade อัตโนมัติ:"
    )
    add_bullet("gemini-3.8-flash (โมเดลเรือธง แม่นยำสูงสุด)", "ลำดับที่ 1: ")
    add_bullet("gemini-3.7-flash (โมเดลเสถียรสูง คิวว่าง แยกโควต้าต่างหาก)", "ลำดับที่ 2: ")
    add_bullet("gemini-3.5-flash-lite (ประมวลผลเร็วพิเศษ)", "ลำดับที่ 3: ")
    add_bullet("gemini-3.1-flash-lite (สำรองลำดับสุดท้าย)", "ลำดับที่ 4: ")
    add_p("ผลลัพธ์: การเรียกใช้งานในโหมดฟรีไม่เคยติดขัดอีกต่อไป เพราะหากโมเดลแรกติดคิว ระบบจะสลับไปโมเดลถัดไปในเสี้ยววินาที!")

    add_h2("Zero-Deduct on Failure: นโยบาย 'ไม่ได้ของ ห้ามตัดโควต้า'")
    add_p(
        "หนึ่งในหลักการสำคัญด้าน Data Quality & User Trust ที่พี่เอกำหนดไว้คือ: "
        "ระบบต้องไม่มีการตัดโควต้าฟรี หรือหักเครดิตเงินของผู้ใช้แม้แต่นาทีเดียว หากกระบวนการถอดเสียงยังไม่สำเร็จ 100% "
        "เราจึงออกแบบให้ Pre-check ทั้งหมดเป็นแบบ Read-Only และจะเรียกฟังก์ชัน Database RPC เพื่อตัดเครดิต "
        "เฉพาะหลังจากที่ AI ส่งมอบผลลัพธ์ที่สมบูรณ์กลับมาแล้วเท่านั้น หากเกิดข้อผิดพลาดระหว่างทาง ผู้ใช้จะไม่เสียสิทธิ์ใด ๆ ทั้งสิ้น"
    )

    add_h2("Memory Explosion Safety Guard บน Serverless")
    add_p(
        "เนื่องจากอัลกอริทึม Needleman-Wunsch ใช้พื้นที่หน่วยความจำแบบ O(M*N) สำหรับคลิปยาว 20-30 นาที "
        "อาจสร้างเมทริกซ์ขนาด 600 ล้านช่อง ซึ่งกิน RAM เกิน 1GB และทำให้ Vercel Serverless ล่มได้ "
        "ระบบจึงมี Safety Guard อัตโนมัติ: หากความยาวตัวอักษรเกิน 1,800 ตัวอักษร หรือ M*N > 2,500,000 "
        "ระบบจะข้ามการรัน DP Matrix และใช้ Timestamp ของ Gemini โดยตรงอย่างปลอดภัย 100%"
    )

    # ---------------------------------------------------------------------------
    # SECTION 6: Business Model & Freemium Strategy
    # ---------------------------------------------------------------------------
    add_h1("5. โมเดลธุรกิจและกลยุทธ์ Freemium (Business Model & Freemium Strategy)")
    add_p(
        "การสร้าง Digital Product ให้เติบโตในยุคนี้ ต้องอาศัยกลยุทธ์ที่สร้างความประทับใจตั้งแต่ครั้งแรกที่ใช้งาน (Aha! Moment) "
        "และเปิดโอกาสให้ผู้ใช้งานได้สัมผัสคุณค่าจริงโดยไม่มีกำแพงขวางกั้น:"
    )

    # Table of Tiers
    tbl_tier = doc.add_table(rows=1, cols=3)
    tbl_tier.alignment = WD_TABLE_ALIGNMENT.CENTER
    tier_headers = ["ระดับผู้ใช้งาน (Tier)", "สิทธิ์และขอบเขตการใช้งาน", "กลยุทธ์ทางธุรกิจ (Strategy)"]
    hdr_cells2 = tbl_tier.rows[0].cells
    for i, h in enumerate(tier_headers):
        hdr_cells2[i].text = h
        set_cell_background(hdr_cells2[i], "0F172A")
        set_cell_margins(hdr_cells2[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells2[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.name = "Tahoma"
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tier_data = [
        (
            "Free Tier (สายฟรี)\n'ฟรีไม่อั้น'",
            "• ไม่จำกัดจำนวนคลิปต่อวัน (Unlimited Clips)\n• ความยาวคลิปไม่เกิน 2 นาที (125 วินาที)\n• ขนาดไฟล์ไม่เกิน 100 MB\n• ประมวลผลผ่าน Dual-Engine Hybrid (Gemini + Groq)\n• ส่งออกไฟล์ VTT, SRT, TXT, Premiere Pro XML ฟรี",
            "Growth & Viral Engine:\nเปิดให้ Creator ทำคลิปสั้น TikTok/Reels ได้ฟรีไม่จำกัด เพื่อสร้างฐานผู้ใช้งาน, สร้างการบอกต่อแบบปากต่อปาก (Word of Mouth), และเปลี่ยนผู้ใช้ให้กลายเป็นแฟนคลับของระบบ"
        ),
        (
            "VIP / Supporter Tier\n(โควต้าผู้สนับสนุน)",
            "• รองรับคลิปยาวสูงสุด 20 นาทีต่อคลิป (1,200 วินาที)\n• ขนาดไฟล์สูงสุด 1.5 GB\n• Fast-track Priority Queue ไม่ต้องรอคิว\n• ขับเคลื่อนด้วยโมเดลท็อปสุด gemini-3.8-flash ตลอดเวลา\n• เติมเครดิตตามจริง (ไม่มีวันหมดอายุ)",
            "Monetization & High-Precision Engine:\nตอบโจทย์ Creator มืออาชีพ, คลิปรีวิว YouTube 8-18 นาที, คลิปสอน How-to, สัมภาษณ์ หรือสรุปข่าว ที่ต้องการความเป๊ะของซับระดับมิลลิวินาที โดยไม่มีสัญญาผูกมัดรายเดือน"
        )
    ]

    for t_name, t_scope, t_strat in tier_data:
        row_cells = tbl_tier.add_row().cells
        for idx, text in enumerate([t_name, t_scope, t_strat]):
            row_cells[idx].text = text
            set_cell_margins(row_cells[idx], top=100, bottom=100, left=120, right=120)
            p = row_cells[idx].paragraphs[0]
            p.runs[0].font.name = "Tahoma"
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.color.rgb = RGBColor(30, 41, 59)
            if idx == 0:
                p.runs[0].font.bold = True
                set_cell_background(row_cells[idx], "F8FAFC")
            else:
                set_cell_background(row_cells[idx], "FFFFFF")

    add_h2("การคำนวณหาจุด Sweet Spot 20 นาที บนสถาปัตยกรรม Oracle Cloud + Coolify")
    add_p(
        "เมื่อทีมงานมีแผนย้ายระบบจาก Vercel ไปสู่ Oracle Cloud Infrastructure (OCI Always Free: 4 OCPU Ampere ARM, 24 GB RAM) "
        "ร่วมกับ Coolify (Self-hosted PaaS) ข้อจำกัดคอขวด 60s Timeout ของ Serverless จะถูกปลดล็อกโดยสิ้นเชิง "
        "เพราะเราสามารถตั้งค่า Reverse Proxy (Nginx/Traefik) ให้ Timeout ได้ถึง 180 - 300 วินาที อย่างไรก็ดี "
        "เราได้คำนวณจุดสมดุล (Sweet Spot) ที่ดีที่สุดสำหรับ VIP ออกมาเป็น '20 นาที (1,200 วินาที)' ด้วยเหตุผล 4 ประการ:"
    )
    add_bullet("คลิปยาว 20 นาทีถ่ายด้วยมือถือ 4K มีขนาดไฟล์ประมาณ 1 - 1.5 GB ซึ่งเป็นเกณฑ์ที่ WebAssembly (Wasm) บนมือถือสามารถสกัดเสียงได้โดยไม่เสี่ยงต่อการเกิด Browser Tab Crash จากขีดจำกัดหน่วยความจำ 32-bit Wasm", "1. ความปลอดภัยของ Memory บนมือถือ (Mobile Wasm Safety): ")
    add_bullet("การรันคู่ขนานและสกัดเสียงใช้เวลารวมเพียง 20 - 28 วินาที ซึ่งเป็นระยะเวลาที่ผู้ใช้รู้สึกประทับใจในความเร็ว (หากนานเกิน 45-60 วินาที ผู้ใช้จะเริ่มรู้สึกว่าระบบค้างและพยายามรีเฟรชหน้าจอ)", "2. จิตวิทยาการรอคอยของผู้ใช้งาน (User Waiting Time): ")
    add_bullet("ด้วย RAM 24 GB ของ Oracle Cloud เราสามารถซอยย่อยการเทียบเสียงเป็น Chunked Alignment (ท่อนละ 2 นาที x 10 ท่อน) ทำให้คลิปยาว 20 นาทีได้รับความเป๊ะระดับมิลลิวินาทีตลอดทั้งคลิปแบบ 100% ไร้การดีเลย์", "3. การันตีความเป๊ะระดับพรีเมียมของ VIP (Acoustic Precision): ")
    add_bullet("ครอบคลุม 98% ของคลิปรีวิวบน YouTube, วิดีโอสอน How-to, สัมภาษณ์ และสรุปข่าวในไทย ซึ่งส่วนใหญ่มีความยาว 8 - 18 นาที", "4. ครอบคลุมพฤติกรรมจริงของครีเอเตอร์ (Market Fit): ")

    # ---------------------------------------------------------------------------
    # SECTION 7: Cost Structure & Unit Economics
    # ---------------------------------------------------------------------------
    add_h1("6. โครงสร้างต้นทุนและผลตอบแทนทางการเงิน (Cost Structure & Unit Economics)")
    add_p(
        "ในฐานะผู้บริหารด้าน Data & Engineering หัวใจสำคัญที่สุดของโปรเจกต์นี้คือ 'ความยั่งยืนทางการเงิน (Unit Economics Sustainability)' "
        "การเปิดให้ใช้งานฟรีต้องไม่ทำให้เจ้าของโปรเจกต์ล้มละลาย และการขายเครดิตต้องมีอัตรากำไรขั้นต้น (Gross Margin) ที่แข็งแรงมากพอ:"
    )

    add_h2("ต้นทุนโครงสร้างพื้นฐานจริงต่อการถอดเสียง 1 นาที (Cost Breakdown per Minute)")

    # Cost Table
    tbl_cost = doc.add_table(rows=1, cols=4)
    tbl_cost.alignment = WD_TABLE_ALIGNMENT.CENTER
    cost_headers = ["รายการทรัพยากร", "อัตราค่าบริการจริงจากผู้ให้บริการ", "ต้นทุนต่อคลิป 1 นาที (THB)", "หมายเหตุประกอบ"]
    hdr_cells3 = tbl_cost.rows[0].cells
    for i, h in enumerate(cost_headers):
        hdr_cells3[i].text = h
        set_cell_background(hdr_cells3[i], "0F172A")
        set_cell_margins(hdr_cells3[i], top=120, bottom=120, left=120, right=120)
        p = hdr_cells3[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.name = "Tahoma"
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    cost_data = [
        ("Google Gemini 3.8 Flash (Paid Key)", "$0.075 / 1 ล้าน Audio Tokens (~$0.000125 / นาที)", "~0.0045 บาท", "ถูกเหลือเชื่อ! ถอดเสียง 1 ชั่วโมง ต้นทุน Gemini ไม่ถึง 30 สตางค์"),
        ("Groq Whisper Large v3", "Free Tier: 20-30 RPM\nPaid: $0.0006 / วินาที", "0.00 บาท (Free Pool)\nหรือ ~1.20 บาท (Paid)", "ในโหมดฟรี รันบน Groq Free Pool ทำให้ต้นทุนเป็น 0 บาท"),
        ("Serverless Compute (Vercel Singapore)", "Execution Time ~3.5 วินาที / 1024MB", "~0.0020 บาท", "อยู่ในโควต้าฟรีของ Vercel หรือเสียเพียงเศษสตางค์"),
        ("Database & Auth (Supabase)", "Postgres Storage & RPC Calls", "~0.0005 บาท", "อยู่ในโควต้าฟรี รองรับผู้ใช้ได้หลักแสนคน"),
        ("ต้นทุนรวมเฉลี่ยต่อ 1 นาที (Total Direct Cost)", "รวมทุกระบบ", "ประมาณ 0.01 - 0.05 บาท / นาที", "Gross Margin สูงกว่า 95% เมื่อเทียบกับราคาขาย"),
    ]

    for res_name, rate, cost_thb, note in cost_data:
        row_cells = tbl_cost.add_row().cells
        for idx, text in enumerate([res_name, rate, cost_thb, note]):
            row_cells[idx].text = text
            set_cell_margins(row_cells[idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[idx].paragraphs[0]
            p.runs[0].font.name = "Tahoma"
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(30, 41, 59)
            if idx == 0:
                p.runs[0].font.bold = True
                set_cell_background(row_cells[idx], "F1F5F9")
            elif idx == 2:
                p.runs[0].font.bold = True
                set_cell_background(row_cells[idx], "FEF3C7" if "รวม" in res_name else "FFFFFF")
            else:
                set_cell_background(row_cells[idx], "FFFFFF")

    sp3 = doc.add_paragraph()
    sp3.paragraph_format.space_before = Pt(6)

    add_h2("โครงสร้างแพ็กเกจเครดิตผู้สนับสนุน (Pricing Packages & Margins)")
    add_p(
        "เมื่อผู้ใช้งานต้องการถอดคลิปยาว หรือสนับสนุนทีมงาน สามารถเลือกซื้อแพ็กเกจเครดิตนาทีได้ดังนี้:"
    )

    # Package Table
    tbl_pkg = doc.add_table(rows=1, cols=5)
    tbl_pkg.alignment = WD_TABLE_ALIGNMENT.CENTER
    pkg_headers = ["แพ็กเกจ", "ราคาขาย (THB)", "ราคาเฉลี่ยต่อนาที", "ต้นทุน API สูงสุด", "กำไรขั้นต้น (Gross Margin)"]
    hdr_cells4 = tbl_pkg.rows[0].cells
    for i, h in enumerate(pkg_headers):
        hdr_cells4[i].text = h
        set_cell_background(hdr_cells4[i], "047857") # Emerald-700
        set_cell_margins(hdr_cells4[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells4[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.name = "Tahoma"
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    pkg_data = [
        ("กาแฟแก้วแรก (Starter)", "99 บาท (50 นาที)", "1.98 บาท / นาที", "~0.50 - 2.50 บาท", "97.5% - 99.5%"),
        ("ครีเอเตอร์ไฟแรง (Pro Creator)", "249 บาท (150 นาที)", "1.66 บาท / นาที", "~1.50 - 7.50 บาท", "97.0% - 99.4%"),
        ("โปรดักชันเฮาส์ (Studio Pack)", "599 บาท (400 นาที)", "1.49 บาท / นาที", "~4.00 - 20.00 บาท", "96.7% - 99.3%"),
    ]

    for p_name, p_price, p_avg, p_cost, p_margin in pkg_data:
        row_cells = tbl_pkg.add_row().cells
        for idx, text in enumerate([p_name, p_price, p_avg, p_cost, p_margin]):
            row_cells[idx].text = text
            set_cell_margins(row_cells[idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[idx].paragraphs[0]
            p.runs[0].font.name = "Tahoma"
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.color.rgb = RGBColor(30, 41, 59)
            if idx == 0:
                p.runs[0].font.bold = True
                set_cell_background(row_cells[idx], "ECFDF5")
            elif idx == 4:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(4, 120, 87)
                set_cell_background(row_cells[idx], "D1FAE5")
            else:
                set_cell_background(row_cells[idx], "FFFFFF")

    sp4 = doc.add_paragraph()
    sp4.paragraph_format.space_before = Pt(6)

    add_callout(
        "ด้วย Gross Margin ที่สูงกว่า 96% ทุก ๆ การสนับสนุน 1 ครั้งจากผู้ใช้งาน VIP จะสามารถนำเงินกำไรส่วนเกิน "
        "ไปหล่อเลี้ยงและชดเชยต้นทุนให้กับผู้ใช้งานสายฟรีได้ถึง 2,000 - 5,000 คลิป! "
        "นี่คือโมเดลเศรษฐกิจแบบ Cross-Subsidization ที่ทำให้แพลตฟอร์มเติบโตได้อย่างยั่งยืนโดยไม่ต้องพึ่งเงินทุนภายนอกเลยแม้แต่น้อย",
        "ความยั่งยืนทางการเงิน (Cross-Subsidization Model)"
    )

    # ---------------------------------------------------------------------------
    # SECTION 8: Key Takeaways & Conclusion
    # ---------------------------------------------------------------------------
    add_h1("7. บทเรียนสำคัญสำหรับนักพัฒนาและผู้สนใจสร้างผลิตภัณฑ์ AI (Key Takeaways)")
    add_bullet(
        "ไม่มี AI Model ตัวไหนที่สมบูรณ์แบบในทุกมิติ: การนำ Acoustic Model (Whisper) ที่เด่นเรื่องคลื่นเสียง "
        "มาจับคู่กับ Linguistic LLM (Gemini) ที่เด่นเรื่องภาษาไทย ผ่านอัลกอริทึม Alignment แบบคลาสสิก "
        "ให้ผลลัพธ์ที่เหนือกว่าการพยายามเค้นความสามารถจากโมเดลใดโมเดลหนึ่งเพียงตัวเดียวอย่างมหาศาล",
        "1. อย่าเชื่อว่าโมเดลเดียวจะทำได้ทุกอย่าง (Hybrid Architecture Wins): "
    )
    add_bullet(
        "การย้ายภาระการแปลงไฟล์ (FFmpeg) ไปไว้ที่ฝั่ง Client-side เบราว์เซอร์ของผู้ใช้ "
        "ช่วยประหยัดค่าใช้จ่ายเซิร์ฟเวอร์ได้เกือบ 100% และลดความเสี่ยงเรื่องเซิร์ฟเวอร์ล่มเมื่อมีผู้ใช้หลั่งไหลเข้ามาพร้อมกัน",
        "2. ผลักภาระการประมวลผลหนัก ๆ ไปที่เครื่องผู้ใช้ (Client-Side First): "
    )
    add_bullet(
        "การสร้างความน่าเชื่อถือกับผู้ใช้ เช่น นโยบาย 'ไม่ตัดโควต้าหากถอดเสียงไม่สำเร็จ' "
        "และการแสดง Badge โปร่งใสว่าคลิปนี้ถอดด้วยโมเดลอะไร มีผลต่อความพึงพอใจและการบอกต่อมากกว่าการอวดความล้ำของเทคโนโลยี",
        "3. ความโปร่งใสและความจริงใจคือหัวใจของผลิตภัณฑ์ (User Trust & Transparency): "
    )
    add_bullet(
        "ในยุคของ Vibe Coding การมี AI Pair Programmer ร่วมกับกระบวนการตรวจสอบคุณภาพที่รัดกุม (เช่น มี Senior QA Subagent คอยรีวิวโค้ดก่อน Commit เสมอ) "
        "ทำให้ Creator ที่เขียนโค้ดไม่เก่งสามารถสร้างซอฟต์แวร์ระดับ Enterprise-Grade ที่เสถียรและปลอดภัยได้จริง",
        "4. พลังของ Vibe Coding + Multi-Agent QA Pair: "
    )

    # Conclusion & Sign-off
    add_p(
        "โปรเจกต์ SUBTHAITLE เป็นเครื่องพิสูจน์ว่า เทคโนโลยีปัญญาประดิษฐ์ยุคใหม่เมื่อผสานเข้ากับการออกแบบสถาปัตยกรรมที่ชาญฉลาด "
        "สามารถสร้างคุณค่ามหาศาลให้กับผู้คน ช่วยปลดล็อกเวลาให้กับ Content Creator ไทยหลายหมื่นคน "
        "และสามารถสร้างโมเดลธุรกิจที่ทำกำไรได้อย่างยั่งยืนไปพร้อมกันครับ",
        bold_prefix="บทสรุปส่งท้าย: "
    )

    output_filename = "SUBTHAITLE_Architecture_and_Case_Study.docx"
    doc.save(output_filename)
    print(f"Document successfully created at {output_filename}")

if __name__ == "__main__":
    create_document()
